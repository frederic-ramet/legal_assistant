# Module pour générer les contrats à partir des templates

import os
import yaml
from datetime import datetime
from pathlib import Path
from docx import Document
from typing import Dict, Any
from .models import Societe


def load_template_config(template_name: str) -> Dict[str, Any]:
    """Charge la configuration d'un template."""
    config_path = Path(f"templates/{template_name}/config.yaml")

    if not config_path.exists():
        raise FileNotFoundError(f"Configuration template introuvable: {config_path}")

    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)


def replace_in_paragraph(paragraph, old_text: str, new_text: str):
    """
    Remplace du texte dans un paragraphe en préservant le formatage.

    Cette fonction gère le cas où le texte est fragmenté en plusieurs runs.
    """
    # Vérifier si le texte à remplacer est dans le paragraphe
    if old_text not in paragraph.text:
        return False

    # Stratégie 1: Essayer le remplacement simple dans les runs
    replaced = False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True

    if replaced:
        return True

    # Stratégie 2: Le texte est fragmenté sur plusieurs runs
    # On reconstruit tout le paragraphe
    full_text = paragraph.text
    if old_text in full_text:
        new_full_text = full_text.replace(old_text, new_text)

        # Supprimer tous les runs existants sauf le premier
        for _ in range(len(paragraph.runs) - 1):
            paragraph._element.remove(paragraph.runs[-1]._element)

        # Mettre le nouveau texte dans le premier run
        if paragraph.runs:
            paragraph.runs[0].text = new_full_text
        else:
            paragraph.add_run(new_full_text)

        return True

    return False


def replace_in_table(table, old_text: str, new_text: str):
    """Remplace du texte dans un tableau."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, old_text, new_text)


def replace_in_document(doc: Document, replacements: Dict[str, str]):
    """Applique tous les remplacements dans le document."""
    # Remplacements dans les paragraphes
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            replace_in_paragraph(paragraph, old_text, new_text)

    # Remplacements dans les tableaux
    for table in doc.tables:
        for old_text, new_text in replacements.items():
            replace_in_table(table, old_text, new_text)


def generate_nda(partie2: Societe, variant: str = "master", output_dir: str = "output") -> str:
    """
    Génère un NDA entre FR Digital (partie 1) et une autre société (partie 2).

    Args:
        partie2: Données de la société partenaire
        variant: Type de NDA (master, dev_plateforme, prestations)
        output_dir: Répertoire de sortie

    Returns:
        Chemin du fichier généré
    """
    print(f"\n🔧 Génération NDA {variant}")
    print(f"   Partie 2: {partie2.raison_sociale}")

    # Charger la config
    config = load_template_config("nda")

    if variant not in config['variants']:
        raise ValueError(f"Variante inconnue: {variant}. Variantes disponibles: {list(config['variants'].keys())}")

    variant_config = config['variants'][variant]
    template_file = variant_config['template']
    format_partie2 = variant_config.get('format_partie2', 'simple')

    # Chemin du template
    template_path = Path(f"templates/nda/examples/{template_file}")

    if not template_path.exists():
        raise FileNotFoundError(f"Template introuvable: {template_path}")

    print(f"   Template: {template_file}")
    print(f"   Format partie 2: {format_partie2}")

    # Charger le document
    doc = Document(template_path)

    # Préparer les remplacements selon le format
    replacements = {}

    if format_partie2 == 'detailed':
        # Format DETAILED (master)
        # Le template contient:
        # "XXXXX , société par actions simplifiée unipersonnelle, dont le siège social
        #  est situé à XXXXX (France), au capital de XXXXX €, inscrit au registre du
        #  commerce de XXXXX sous le numéro d'inscription XXXXX , dûment représenté par XXXXX ,"

        # Note: Il y a un espace avant la virgule après le premier XXXXX dans le template
        replacements = {
            # Premier XXXXX (raison sociale) - avec espace avant la virgule
            "XXXXX , société par actions simplifiée unipersonnelle":
                f"{partie2.raison_sociale}, {partie2.forme_juridique.lower()}",
            # Les autres XXXXX dans leur contexte
            "situé à XXXXX (France)": f"situé à {partie2.adresse}",
            "capital de XXXXX €": f"capital de {partie2.capital}",
            "commerce de XXXXX sous": f"commerce de {partie2.ville_rcs} sous",
            "numéro d'inscription XXXXX ": f"numéro d'inscription {partie2.siren} ",
            "représenté par XXXXX ,": f"représenté par {partie2.representant_nom},",
        }

    else:
        # Format SIMPLE (dev_plateforme, prestations)
        # XXXXXXX, dont le siège social est situé àXXXXXXX, inscrit au registre du
        # commerce de XXX sous le numéro d'inscription XXXXXXX, dûment représenté par XXXXX,

        replacements = {
            "XXXXXXX, dont le siège social est situé àXXXXXXX":
                f"{partie2.raison_sociale}, dont le siège social est situé à {partie2.adresse}",
            "commerce de XXX sous": f"commerce de {partie2.ville_rcs} sous",
            f"inscription XXXXXXX,": f"inscription {partie2.siren},",
            f"par XXXXX,": f"par {partie2.representant_nom},",
        }

    # Appliquer les remplacements
    replace_in_document(doc, replacements)

    # Signatures (tableau)
    signature_replacements = {
        "Nom :": f"Nom : {partie2.representant_nom}",
        "Titre :": f"Titre : {partie2.representant_fonction}",
    }

    for table in doc.tables:
        # Chercher le tableau de signatures
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if "LE PARTENAIRE" in cell_text or "Nom :" in cell_text:
                    for old, new in signature_replacements.items():
                        for paragraph in cell.paragraphs:
                            if old in paragraph.text:
                                replace_in_paragraph(paragraph, old, new)

    # Créer le répertoire de sortie
    os.makedirs(output_dir, exist_ok=True)

    # Nom du fichier de sortie
    date_str = datetime.now().strftime("%d-%m-%Y")
    partie1_clean = "FRDigital"
    partie2_clean = partie2.raison_sociale.replace(' ', '').replace('/', '')[:20]

    output_filename = f"NDA_{variant}_{partie1_clean}_{partie2_clean}_{date_str}.docx"
    output_path = Path(output_dir) / output_filename

    # Sauvegarder
    doc.save(output_path)

    print(f"✅ NDA généré: {output_path}")

    return str(output_path)


if __name__ == "__main__":
    # Test
    from .scraper import scrape_pappers

    print("Test generator NDA\n")

    # Scraper les données
    print("=" * 60)
    nexans = scrape_pappers("393525852")

    print("\n" + "=" * 60)
    # Générer un NDA master
    output_file = generate_nda(nexans, variant="master")

    print(f"\n📄 Fichier généré: {output_file}")
